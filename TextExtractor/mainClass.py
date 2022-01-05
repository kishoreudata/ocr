# -*- coding: utf-8 -*-
"""
Created on Tue Aug 10 16:22:28 2021
@author: Sunil
"""
import sys
from PyQt5 import QtWidgets, QtCore
from PyQt5.QtWidgets import QFileDialog
import os
import time
import docx
import json
from PyQt5.QtGui import QTextCursor, QIcon, QPixmap
import cv2
import re
import enchant
from pdf2image import convert_from_path

from mainClassGui import Ui_MainWindow

from PIL import Image
from pytesseract import pytesseract
from matToQimage import toQImage

# Defining paths to tesseract.exe
path_to_tesseract = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

class MainClassStart(Ui_MainWindow):
     
    def __init__(self, Form):
        print("\n --------Start init Program--------")
        
        Ui_MainWindow.__init__(self)
        self.setupUi(Form)
        
        self.pushButton.clicked.connect(self.selectFileButton)
        self.pushButton_2.clicked.connect(self.SaveButton)
        self.pushButton_3.clicked.connect(self.ExtractButton)
        self.dict = enchant.Dict("en_US") 
        self.lang='Devanagari'
        self.isEn=True
        print("\n --------End init Program--------")
        
    def selectFileButton(self):
        print("\n ------- selectFileButton Started -------")
        
        global fileName 
        global imagefile
        options=QFileDialog.Options()
        fileName, _=QFileDialog.getOpenFileName(None ,"QFileDialog.getOpenFileName()","", "All Files (*);; Python Files (*.py)", options=options)
        #fileName = QFileDialog.getExistingDirectory(
            #None, 'Select a directory', '')
        #if fileName:
        print(fileName)
        self.label_2.setText(fileName)
        
        w = self.label_4.width()
        h = self.label_4.height()
        
        file_name, file_extension = os.path.splitext(fileName)
        
        if(file_extension=='.pdf'):
            print('Pdf File is Selected')
            images = convert_from_path(fileName, 500,poppler_path=r'C:\Program Files\poppler-0.67.0\bin')
            for i in range(len(images)):
                print(images[i])
                images[i].save('savedimg'+ str(i) +'.jpg', 'JPEG')
                time.sleep(1)
                imagefile = cv2.imread('savedimg0.jpg')
        else:
            imagefile = cv2.imread(fileName)
        
        img = toQImage(imagefile)
        #print(img1)        
        #w = self.label_4.width()
        #h = self.label_4.height()
        
        pixmap = QPixmap(img)
        pixmap1 = pixmap.scaled(w, h, transformMode=QtCore.Qt.SmoothTransformation)
        self.label_4.setPixmap(pixmap1)
        
        print("\n ------- selectFileButton Completed -------")
        
    def sanitize_str(self, s):
        control_chars = "\x00-\x1f\x7f-\x9f"
        control_char_re = re.compile("[%s]" % control_chars)
        return control_char_re.sub("", s)
    
    def SaveButton(self):
        print("\n ------- SaveButton Started -------")
        linestr = ''
        doc = docx.Document()
        
        #mytext = self.textEdit.toHtml()
        mytext = self.textEdit.toPlainText()
        print(type(mytext))
        
        for line in mytext:
            #print(line)
            if(line=='\n'):
                doc.add_paragraph(self.sanitize_str(linestr))
                linestr = ''
            elif(line=='\n' and linestr==''):
                continue
            else:
                linestr = linestr + line
        
        doc.save('Transcription.docx')
        time.sleep(1)
        options=QFileDialog.Options()
        fileName1, _=QFileDialog.getSaveFileName(None, "QFileDialog.getSaveFileName()","untitled", "All Files (*); Docx Files (*.docx)", options=options)
        doc.add_paragraph("")
        doc.save(fileName1)
        
        print("\n ------- SaveButton Completed -------")   

    def ExtractButton(self):
        print("\n --------ExtractButton Started--------")
        
        # Providing the tesseract executable
        # location to pytesseract library
        enhindw=""
        pytesseract.tesseract_cmd = path_to_tesseract
        
        # Passing the image object to image_to_string() function
        # This function will extract the text from the image
        #text = pytesseract.image_to_string(imagefile, lang='eng+hin+ind+dan+urd')
        text = pytesseract.image_to_string(imagefile,self.lang )
        words = text.split()
        lwrd=[]
        for wrd in words:
            if not self.isEnglishWord(wrd):
                lwrd.append(wrd)
        enhindw.join(lwrd) # outputs nondictionary english as well as hindi words
        self.textEdit.setText(enhindw)
        
        print("\n --------ExtractButton Completed--------")
        
    def ClearButton(self):
        print("\n ------- clearButton Started -------")
        
        self.textEdit.setText("")
        
        print("\n ------- clearButton completed -------")
    
    def isEnglishWord(self,word):
        self.isEn= self.dict.check(word)
        return self.isEn
    
    def extractHindi(self,wrd):
        hindw=''
        pattern = re.compile("^[a-zA-Z@_!#$%^&*()<>?\|}{~:.;]")
        for ch in wrd:
            if not pattern.match(ch):
               hindw+=ch#outputs only hindi words
        return hindw

if __name__ == "__main__":
    
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    mainObject = MainClassStart(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())

#
# End Of Program
